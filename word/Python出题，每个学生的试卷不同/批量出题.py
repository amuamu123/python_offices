import random
#不重复随机整数生成函数
def Random_num(num_max, num_qty):
    '''
    num_max:最大数
    num_qty:生成随机数的个数
    '''
    num_list = [] #储存生成的随机数
    while len(num_list) < num_qty: #控制随机数的个数
        num = random.randint(2,num_max) #设定在此范围内取数
        if num in num_list: #判断随机数是否重复
            continue #若重复，则重新生成
        else:
            num_list.append(num) #将不重复的随机数放入列表
    return num_list #生成完成后返回随机数列表

#定义函数，按随机数在题库中抽取对应编号的题目
from openpyxl import load_workbook
def Question(que_type,numbers):
    '''
    que_type:试题类型（"单选题","多选题","判断题","填空题"）
    numbers:需要抽取的试题编号
    '''
    questions = [] #储存抽取的题目
    wb = load_workbook("题库.xlsx") #载入题库
    if que_type == "单选题":
        ws = wb[que_type]        
        for i in numbers: #按随机生成的编号抽题
            question = ws["B"+str(i)].value #问题在B列
            answerA = "A：\t" + str(ws["C"+str(i)].value) #选项A在C列，"\t"相当于按一下tab键，在字符间产生间隔
            answerB = "B：\t" + str(ws["D"+str(i)].value) #选项B在D列
            answerC = "C：\t" + str(ws["E"+str(i)].value) #选项C在E列
            answerD = "D：\t" + str(ws["F"+str(i)].value) #选项D在F列
            right_answer  = ws["G"+str(i)].value #正确答案在G列
            single_question = [question, answerA, answerB, answerC, answerD, right_answer] #每行的数据存入列表
            questions.append(single_question) #每个题目的数据存入总列表
    elif que_type == "多选题":
        ws = wb[que_type]        
        for i in numbers:
            question = ws["B"+str(i)].value            
            answerA = "A：\t" + str(ws["C"+str(i)].value)
            answerB = "B：\t" + str(ws["D"+str(i)].value)
            answerC = "C：\t" + str(ws["E"+str(i)].value)
            answerD = "D：\t" + str(ws["F"+str(i)].value)
            right_answer  = ws["H"+str(i)].value
            single_question = [question, answerA, answerB, answerC, answerD, right_answer]
            if ws["G"+str(i)].value: #有些题有E选项，有些没有，因此需要判断一下是否有E选项
                answerE = "E：\t" + str(ws["G"+str(i)].value)
                single_question.insert(-1,answerE) #将E选项插入到答案前面，保持答案是最后一个元素
            questions.append(single_question)
    else: #判断题和填空题，内容只取题干和答案
        ws = wb[que_type]        
        for i in numbers:
            question = ws["B"+str(i)].value
            right_answer  = ws["C"+str(i)].value
            single_question = [question, right_answer]
            questions.append(single_question)
            
    return questions

#写入考试题到word文件
from docx import Document
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体
def To_word(number,questions_data):
    doc = Document("试题-模板.docx")

    #写入单选题
    title1 = "一、单项选择题（共40题，每题1分）"
    p = doc.add_paragraph() #插入段落
    r = p.add_run(title1) #插入文字块
    r.bold = True #字体加粗
    r.font.size = Pt(12) #字号设为12磅
    for index, i in enumerate(questions_data["单选题"],start = 1): #给题目从1开始编号
        doc.add_paragraph(f"{index}. {i[0]}") #题干部分在单独一段
        doc.add_paragraph(f"\t{i[1]}\t\t{i[2]}") #选项A和选项B在同一段落
        doc.add_paragraph(f"\t{i[3]}\t\t{i[4]}") #选项C和选项D在同一段落

    #写入多选题
    title2 = "二、多项选择题（共20题，每题2分）"
    p = doc.add_paragraph()
    r = p.add_run(title2)
    r.bold = True
    r.font.size = Pt(12)
    for index, i in enumerate(questions_data["多选题"],start = 1):
        doc.add_paragraph(f"{index}. {i[0]}")
        doc.add_paragraph(f"\t{i[1]}\t\t{i[2]}")
        doc.add_paragraph(f"\t{i[3]}\t\t{i[4]}")
        if len(i) ==7: #判断是否有E选项，若有，则新建一段落写入
            doc.add_paragraph(f"\t{i[5]}")

    #写入判断题
    title3 = "三、判断题（共10题，每题1分）"
    p = doc.add_paragraph()
    r = p.add_run(title3)
    r.bold = True
    r.font.size = Pt(12)
    for index, i in enumerate(questions_data["判断题"],start = 1):
        doc.add_paragraph(f"\t{index}. {i[0]}")

    #写入填空题
    title4 = "四、填空题（共10题，每题1分）"
    p = doc.add_paragraph()
    r = p.add_run(title4)
    r.bold = True
    r.font.size = Pt(12)
    for index, i in enumerate(questions_data["填空题"],start = 1):
        doc.add_paragraph(f"\t{index}. {i[0]}")

    doc.save(f"试卷及答案\\考试题{number}.docx")

#写入答案
from docx import Document
from docx.shared import Pt #用于设定字体大小（磅值）
from docx.oxml.ns import qn #用于应用中文字体

def Answer(number,questions_data):
    doc = Document()
    #全局字体设为“宋体”
    doc.styles['Normal'].font.name=u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    title = "计算机系2020第二学期期末考试题(答案)"
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    #写入单选题答案
    title1 = "一、单项选择题答案（共40题，每题1分）"
    p = doc.add_paragraph()
    r = p.add_run(title1)
    r.bold = True
    r.font.size = Pt(12)
    
    p = doc.add_paragraph()
    for index, i in enumerate(questions_data["单选题"],start = 1):
        p.add_run(f"{index}. {i[-1]}\t")
        if index%10 == 0: #每段只显示10个答案
            p = doc.add_paragraph() #满10个，则新建段落

    #写入多选题答案
    title2 = "二、多项选择题答案（共20题，每题2分）"
    p = doc.add_paragraph()
    r = p.add_run(title2)
    r.bold = True
    r.font.size = Pt(12)
    p = doc.add_paragraph()
    for index, i in enumerate(questions_data["多选题"],start = 1):
        p.add_run(f"{index}. {i[-1]}\t")
        if index%3 == 0: 
            p = doc.add_paragraph() 
        
    #写入判断题答案
    title3 = "三、判断题答案（共10题，每题1分）"
    p = doc.add_paragraph()
    r = p.add_run(title3)
    r.bold = True
    r.font.size = Pt(12)
    p = doc.add_paragraph()
    for index, i in enumerate(questions_data["判断题"],start = 1):
        p.add_run(f"{index}. {i[-1]}\t")
        if index%5 == 0: #每段只显示5个答案
            p = doc.add_paragraph() #满5个，则新建段落

    #写入填空题
    title4 = "四、填空题答案（共10题，每题1分）"
    p = doc.add_paragraph()
    r = p.add_run(title4)
    r.bold = True
    r.font.size = Pt(12)
    p = doc.add_paragraph()
    for index, i in enumerate(questions_data["填空题"],start = 1):
        p.add_run(f"{index}. {i[-1]}\t\t")
        if index%2 == 0: #每段只显示2个答案
            p = doc.add_paragraph() #满2个，则新建段落
            
    doc.save(f"试卷及答案\\考试题{number}答案.docx")

#主函数
for number in range(1,21): #不同的试卷数量，此处为20套
    #生成随机题目编号    
    num_single_choice = Random_num(566,40)
    num_mult_choice = Random_num(196,20)
    num_judgment = Random_num(418,10)
    num_completion = Random_num(190,10)
    #将生成的编号存入字典`question_num`
    question_num = {"单选题号":num_single_choice,
           "多选题号":num_mult_choice,
            "判断题号":num_judgment,
            "填空题号":num_completion
           }
    #根据随机生成的题目编号去题库选题，并存入`questions_data`
    questions_data = {
        "单选题":Question("单选题",question_num["单选题号"]),
        "多选题":Question("多选题",question_num["多选题号"]),
        "判断题":Question("判断题",question_num["判断题号"]),
        "填空题":Question("填空题",question_num["填空题号"])
    }
    #将试题写入word文档，并保存
    To_word(number,questions_data)
    #将试题答案写入word文档，并保存
    Answer(number,questions_data)
    print(f"试卷{number}及答案完成！")