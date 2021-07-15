import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# 设置目标文件夹路径
path = "涨薪通告-练习/"

# 获取目标文件夹下的所有文件名
file_list = os.listdir(path)

for file in file_list:
    # 拼接文件路径
    file_path = path + file
    
    # 打开 Word 文件
    doc = Document(file_path)

    # 添加 Paragraph 对象 para_1
    para_1 = doc.add_paragraph('盖章: ') # 添加文字
    # 添加 Run 对象 run_stamp
    run_stamp = para_1.add_run()
    run_stamp.add_picture("Shining.png") # 添加图片

    # 添加 Paragraph 对象 para_2
    para_2 = doc.add_paragraph()
    # 设置对齐方式
    para_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # 添加 Run 对象 run_comp
    run_comp = para_2.add_run('闪光科技金融公司(Shining Fintech Company)')
    # 设置字体
    run_comp.font.size = Pt(14) # 字体大小
    run_comp.font.bold = True # 字体加粗

    # 保存文件
    doc.save(file_path)