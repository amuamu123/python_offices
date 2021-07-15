# 案例 1：批量生成法务函
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

# 打开【封号名单.xlsx】工作簿，获取【封号人员】工作表
title_wb = load_workbook('../工作/封号名单.xlsx')
title_sheet = title_wb.active

# 遍历【封号人员】工作表中的数据
for row in title_sheet.iter_rows(min_row=2, values_only=True):
    # 获取“名字”
    name = row[0]
    # 获取“微信号”
    wxid = row[1]

    # 打开 Word 文件
    doc = Document('../工作/法务函模板.docx')

    # 获取第六个 Paragraph 对象
    para = doc.paragraphs[5]

    # 添加封号人员名字
    run_name = para.add_run(name)
    # 字体加粗
    run_name.font.bold = True
    # 字体加下划线
    run_name.font.underline = True
    # 设置字体大小为 14 pt
    run_name.font.size = Pt(14)

    # 添加封号人员微信号
    run_wxid = para.add_run(' 同学（WeChat ID: {}）'.format(wxid))
    # 设置字体大小为 14 pt
    run_wxid.font.size = Pt(14)

    # 按封号人员名字命名并保存文件到【法务函文件】文件夹
    doc.save('../工作/法务函文件/法务函-{}.docx'.format(name))